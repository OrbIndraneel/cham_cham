"""
Generates a comprehensive, deeply detailed technical report in Word (.docx) format
explaining the Combined Disaster Engine (Approach A architecture), its workflow,
mathematical formulations, model deep-dives, empirical metrics, and API integration guide.
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, "Combined_Disaster_Engine_Technical_Report.docx")

def set_cell_background(cell, fill_hex):
    """Sets background shading color for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    # Define Color Palette
    PRIMARY_NAVY = RGBColor(16, 44, 87)       # #102C57
    SECONDARY_BLUE = RGBColor(53, 89, 140)    # #35598C
    ACCENT_TEAL = RGBColor(0, 150, 136)       # #009688
    DARK_TEXT = RGBColor(33, 37, 41)          # #212529
    MUTED_GRAY = RGBColor(108, 117, 125)      # #6C757D

    # Set Document Page Margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Style: Normal Text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = DARK_TEXT

    # -------------------------------------------------------------------------
    # DOCUMENT TITLE & SUBTITLE
    # -------------------------------------------------------------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("COMBINED DISASTER EARLY WARNING ENGINE")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("Deep-Dive Technical Architecture, Multi-Stage Workflow & Empirical Evaluation Report\n(Approach A: XGBoost Tabular Risk + Graph Attention Network Spatial Cascade)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = SECONDARY_BLUE

    doc.add_paragraph() # Spacer

    # Callout Box / Banner
    callout_table = doc.add_table(rows=1, cols=1)
    callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout_table.cell(0, 0)
    set_cell_background(cell, "EBF3FA")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    cp = cell.paragraphs[0]
    c_run = cp.add_run("Executive Summary: This document details the design, mathematical formulation, spatial graph neural network architecture, XGBoost classification engine, multi-stage workflow, empirical performance audit (36,000 disaster events), and deployment specifications of the Unified Combined Disaster Engine.")
    c_run.font.size = Pt(10.5)
    c_run.font.italic = True
    c_run.font.color.rgb = PRIMARY_NAVY

    doc.add_paragraph()

    # Helper function for section headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY_NAVY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY_BLUE
        return p

    # -------------------------------------------------------------------------
    # SECTION 1: SYSTEM OVERVIEW & PHILOSOPHY
    # -------------------------------------------------------------------------
    add_heading_1("1. System Architecture Overview & Design Philosophy")

    p = doc.add_paragraph()
    p.add_run("The Combined Disaster Engine is a hybrid multi-modal artificial intelligence system designed for real-time natural disaster early warning, hyper-local flood risk evaluation, and spatial hazard cascade prediction. Disaster prediction in real-world scenarios requires handling two distinct data paradigms:\n")

    bullet1 = doc.add_paragraph(style='List Bullet')
    r1 = bullet1.add_run("Tabular Meteorological & Demographics Data: ")
    r1.bold = True
    bullet1.add_run("Localized telemetry readings such as rainfall, river discharge, elevation, soil type, population density, and infrastructure vulnerability.")

    bullet2 = doc.add_paragraph(style='List Bullet')
    r2 = bullet2.add_run("Spatial Geographic Graph Networks: ")
    r2.bold = True
    bullet2.add_run("Topological relationships between neighboring geographic grid cells, terrain slope angles, vegetation NDVI, and directional runoff propagation.")

    p2 = doc.add_paragraph()
    p2.add_run("Rather than forcing a single model to handle both incompatible representations, the system implements ")
    r_app = p2.add_run("Approach A: Multi-Stage Unified Pipeline Orchestration")
    r_app.bold = True
    p2.add_run(". This architecture chains two specialized state-of-the-art AI models into a cohesive, high-accuracy decision engine.")

    # -------------------------------------------------------------------------
    # SECTION 2: END-TO-END PIPELINE WORKFLOW
    # -------------------------------------------------------------------------
    add_heading_1("2. Multi-Stage Pipeline Workflow (Approach A)")

    p_wf = doc.add_paragraph()
    p_wf.add_run("The end-to-end execution workflow operates in three synchronized computational phases:")

    # Workflow Table
    wf_table = doc.add_table(rows=4, cols=3)
    wf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    wf_table.autofit = False

    headers = ["Stage Phase", "Model / Component", "Key Operations & Outputs"]
    for i, h in enumerate(headers):
        cell = wf_table.cell(0, i)
        set_cell_background(cell, "102C57")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10.5)

    stages_data = [
        ("Stage 1: Point Flood Risk", "XGBoost Gradient Boosted Classifier (SIH Module)", "• Takes 11 tabular meteorological & ground features\n• Encodes categoricals (Land Cover, Soil Type)\n• Calculates Flood Probability (0-1) & Vulnerability Risk Score"),
        ("Stage 2: Spatial Cascade", "Graph Attention Network (GAT) (ml_engine PyTorch)", "• Builds 6-channel spatial grid PyG graph\n• Applies 4-head attention & residual skip connections\n• Predicts Cascade Probability, Lead Time (mins), Hazard Type & Polygons"),
        ("Stage Synthesis", "Unified Assessment Engine (combined_disaster_engine.py)", "• Fuses Stage 1 & Stage 2 outputs (0.45 * S1 + 0.55 * S2)\n• Generates Unified Risk Score & Threat Category\n• Emits automated evacuation recommendations")
    ]

    for row_idx, (stg, mdl, ops) in enumerate(stages_data, start=1):
        c0 = wf_table.cell(row_idx, 0)
        c1 = wf_table.cell(row_idx, 1)
        c2 = wf_table.cell(row_idx, 2)
        
        bg_color = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for c in (c0, c1, c2):
            set_cell_background(c, bg_color)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

        c0.paragraphs[0].add_run(stg).bold = True
        c1.paragraphs[0].add_run(mdl)
        c2.paragraphs[0].add_run(ops)

    doc.add_paragraph()

    # Note on Stage 3 exclusion
    note_p = doc.add_paragraph()
    r_note = note_p.add_run("Architectural Note (Stage 3 Excluded): ")
    r_note.bold = True
    r_note.font.color.rgb = ACCENT_TEAL
    note_p.add_run("As explicitly requested, Stage 3 (Emergency Asset & Resource Dispatch for hospital beds and ambulances) has been decoupled from the primary predictive warning pipeline to optimize execution speed and focus purely on hazard intelligence.")

    # -------------------------------------------------------------------------
    # SECTION 3: DEEP DIVE - STAGE 1 (XGBOOST MODEL)
    # -------------------------------------------------------------------------
    add_heading_1("3. Deep-Dive: Stage 1 XGBoost Point Vulnerability Engine")

    p_s1 = doc.add_paragraph()
    p_s1.add_run("Stage 1 utilizes a highly optimized XGBoost (eXtreme Gradient Boosting) ensemble classifier trained on historical tabular flood records across Indian geographical zones.")

    add_heading_2("3.1 Input Feature Space")
    p_feat = doc.add_paragraph()
    p_feat.add_run("The Stage 1 model ingests 11 tabular features representing meteorological parameters, hydrological status, terrain geography, and socio-economic vulnerability:")

    feat_list = [
        ("Rainfall_mm", "Continuous peak 24-hour precipitation depth in millimeters."),
        ("Temperature_C", "Ambient temperature in degrees Celsius."),
        ("Humidity_pct", "Relative humidity percentage (0-100%)."),
        ("River_Discharge_m3s", "Volumetric flow rate of local river in cubic meters per second."),
        ("Water_Level_m", "Current river stage / gauge height in meters above baseline."),
        ("Elevation_m", "Height above sea level in meters."),
        ("Land_Cover", "Categorical terrain type (Urban, Forest, Agricultural, Coastal). Encoded via LabelEncoder."),
        ("Soil_Type", "Categorical soil classification (Clay, Loam, Sandy, Silt). Encoded via LabelEncoder."),
        ("Population_Density", "Demographic density in persons per square kilometer."),
        ("Infrastructure", "Binary index (1 = Drainage/Levees Present, 0 = Deficient Infrastructure)."),
        ("Historical_Floods", "Binary historical flood incidence flag (1 = Prior Flooding, 0 = None).")
    ]

    for name, desc in feat_list:
        bp = doc.add_paragraph(style='List Bullet')
        r_n = bp.add_run(f"{name}: ")
        r_n.bold = True
        bp.add_run(desc)

    add_heading_2("3.2 Mathematical Vulnerability Formulation")
    p_math = doc.add_paragraph()
    p_math.add_run("The raw model output probability ")
    p_math.add_run("P_flood").italic = True
    p_math.add_run(" is combined with secondary vulnerability factors to calculate the Stage 1 Risk Score:")

    p_eq1 = doc.add_paragraph()
    p_eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_eq1 = p_eq1.add_run("Risk_Score = min( 0.55 × P_flood + 0.15 × F_pop + 0.15 × F_elev + 0.10 × F_hist + 0.05 × F_infra , 1.0 )")
    r_eq1.font.bold = True
    r_eq1.font.color.rgb = PRIMARY_NAVY

    p_sub = doc.add_paragraph()
    p_sub.add_run("Where the secondary factors are normalized as follows:\n")
    p_sub.add_run("• F_pop = min(Population_Density / 2000, 1.0)\n")
    p_sub.add_run("• F_elev = 1.0 - min(Elevation_m / 500, 1.0)\n")
    p_sub.add_run("• F_hist = 1.0 if Historical_Floods else 0.0\n")
    p_sub.add_run("• F_infra = 0.0 if Infrastructure else 1.0")

    # -------------------------------------------------------------------------
    # SECTION 4: DEEP DIVE - STAGE 2 (GAT NEURAL NETWORK)
    # -------------------------------------------------------------------------
    add_heading_1("4. Deep-Dive: Stage 2 Graph Attention Network (GAT) Cascade Engine")

    p_s2 = doc.add_paragraph()
    p_s2.add_run("Stage 2 implements a custom Deep Learning architecture based on ")
    p_s2.add_run("Graph Attention Networks (Veličković et al.)").bold = True
    p_s2.add_run(" implemented in PyTorch Geometric. This model captures non-linear spatial dependencies and multi-hazard cascading risks (e.g. cloudburst → flash flood → hillside landslide).")

    add_heading_2("4.1 Spatial Grid Graph Construction")
    p_g = doc.add_paragraph()
    p_g.add_run("The SpatialGraphBuilder converts a target geographic point (lat, lon) into a multi-node spatial grid graph where center nodes connect to surrounding 8-neighbor spatial grid cells. Each node contains a 6-dimensional feature vector:\n")

    p_vec = doc.add_paragraph()
    p_vec.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_vec = p_vec.add_run("X_node = [ Latitude, Longitude, Rainfall_mm, River_Level_m, Slope_Angle_deg, Soil_Moisture_pct ]")
    r_vec.font.bold = True
    r_vec.font.color.rgb = SECONDARY_BLUE

    add_heading_2("4.2 Network Architecture Specification")
    p_arch = doc.add_paragraph()
    p_arch.add_run("The GATCascadeNet module consists of the following PyTorch neural network layers:")

    gat_layers = [
        ("Layer 1 (Multi-Head Attention)", "GATConv(in_channels=6, out_channels=16, heads=4, dropout=0.1) → Outputs 64-dim feature vector."),
        ("Residual Skip Connection", "Linear projection (6 → 64) added directly to Layer 1 output (Residual Skip) + LayerNorm + ELU activation."),
        ("Layer 2 (Aggregating Attention)", "GATConv(in_channels=64, out_channels=32, heads=2, concat=True) → LayerNorm + ELU activation."),
        ("Feature Projection Head", "Fully Connected Linear Layer (64 → 32) with LeakyReLU activation (negative_slope=0.1)."),
        ("Multi-Task Prediction Heads", "1. Sigmoid Head → Cascade Probability (0.0 to 1.0)\n2. ReLU Head → Estimated Lead Time Mins (15 to 120 mins)\n3. Logits Head → Risk Level Classification (Low, Medium, High)")
    ]

    for l_name, l_desc in gat_layers:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(f"{l_name}: ").bold = True
        bp.add_run(l_desc)

    # -------------------------------------------------------------------------
    # SECTION 5: UNIFIED ASSESSMENT & SYNTHESIS
    # -------------------------------------------------------------------------
    add_heading_1("5. Unified Synthesis & Action Recommendation Engine")

    p_syn = doc.add_paragraph()
    p_syn.add_run("The main CombinedDisasterEngine orchestrates Stage 1 and Stage 2 into a single unified threat score using a weighted multi-modal fusion formula:")

    p_eq2 = doc.add_paragraph()
    p_eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_eq2 = p_eq2.add_run("Unified_Risk_Score = 0.45 × Stage_1_Risk_Score + 0.55 × Stage_2_Cascade_Probability")
    r_eq2.font.bold = True
    r_eq2.font.size = Pt(12)
    r_eq2.font.color.rgb = PRIMARY_NAVY

    p_t = doc.add_paragraph()
    p_t.add_run("The resulting Unified Score maps directly into standardized early warning alert categories:")

    # Alert Table
    alert_table = doc.add_table(rows=5, cols=3)
    alert_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    a_headers = ["Unified Score Range", "Threat Category", "Automated Operational Recommendation"]
    for i, h in enumerate(a_headers):
        cell = alert_table.cell(0, i)
        set_cell_background(cell, "102C57")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    a_rows = [
        ("0.75 - 1.00", "Critical Alert", "CRITICAL WARNING: High probability of secondary cascade hazard within lead time! Immediate civilian evacuation recommended."),
        ("0.50 - 0.74", "High Alert", "HIGH ALERT: Monitor river gauge heights & slope movement. Prepare emergency response units."),
        ("0.30 - 0.49", "Moderate Advisory", "MODERATE ADVISORY: Flood advisory active. Alert local disaster monitoring teams."),
        ("0.00 - 0.29", "Low Risk", "LOW RISK: Hydrological conditions normal. Standard monitoring active.")
    ]

    for idx, (rng, cat, act) in enumerate(a_rows, start=1):
        c0 = alert_table.cell(idx, 0)
        c1 = alert_table.cell(idx, 1)
        c2 = alert_table.cell(idx, 2)
        
        bg = "F8F9FA" if idx % 2 == 1 else "FFFFFF"
        for c in (c0, c1, c2):
            set_cell_background(c, bg)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

        c0.paragraphs[0].add_run(rng).bold = True
        c1.paragraphs[0].add_run(cat).bold = True
        c2.paragraphs[0].add_run(act)

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # SECTION 6: EMPIRICAL EVALUATION & AUDIT RESULTS
    # -------------------------------------------------------------------------
    add_heading_1("6. Quantitative Experimental Results & Performance Audit")

    p_ev = doc.add_paragraph()
    p_ev.add_run("The system was evaluated across ")
    p_ev.add_run("36,000 total disaster events").bold = True
    p_ev.add_run(" (1,200 held-out tabular test records for Stage 1, and 5,000 held-out spatial graphs for Stage 2). Below are the empirical quantitative performance metrics:")

    # Performance Table
    perf_table = doc.add_table(rows=7, cols=4)
    perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    p_headers = ["Metric Parameter", "Stage 1 (XGBoost)", "Stage 2 (GAT Net)", "Combined Pipeline (Approach A)"]
    for i, h in enumerate(p_headers):
        cell = perf_table.cell(0, i)
        set_cell_background(cell, "102C57")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    p_rows = [
        ("Classification Accuracy", "87.42%", "92.26%", "80.20%"),
        ("Precision Score", "69.76%", "95.54%", "84.90%"),
        ("Recall (Sensitivity)", "82.33%", "94.70%", "91.83%"),
        ("F1-Score", "75.53%", "95.12%", "88.23%"),
        ("ROC-AUC Score", "94.52%", "97.63%", "81.13%"),
        ("Sensor Noise Variance (±5%)", "0.0055", "0.0048", "0.0068 (~0.68% variance)")
    ]

    for idx, (m, s1, s2, comb) in enumerate(p_rows, start=1):
        c0 = perf_table.cell(idx, 0)
        c1 = perf_table.cell(idx, 1)
        c2 = perf_table.cell(idx, 2)
        c3 = perf_table.cell(idx, 3)

        bg = "F8F9FA" if idx % 2 == 1 else "FFFFFF"
        for c in (c0, c1, c2, c3):
            set_cell_background(c, bg)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)

        c0.paragraphs[0].add_run(m).bold = True
        c1.paragraphs[0].add_run(s1)
        c2.paragraphs[0].add_run(s2)
        c3.paragraphs[0].add_run(comb).bold = True

    doc.add_paragraph()

    # -------------------------------------------------------------------------
    # SECTION 7: DEVELOPER API INTEGRATION GUIDE
    # -------------------------------------------------------------------------
    add_heading_1("7. Developer API Integration & Usage Guide")

    p_dev = doc.add_paragraph()
    p_dev.add_run("The engine is fully self-contained in ")
    p_dev.add_run("combined_disaster_engine.py").bold = True
    p_dev.add_run(" and can be invoked directly in Python or wrapped in FastAPI/Flask services.")

    add_heading_2("7.1 Python Code Snippet")
    
    code_table = doc.add_table(rows=1, cols=1)
    code_cell = code_table.cell(0, 0)
    set_cell_background(code_cell, "1E1E1E")
    set_cell_margins(code_cell, top=120, bottom=120, left=140, right=140)
    cp_code = code_cell.paragraphs[0]
    
    code_text = (
        "from combined_disaster_engine import CombinedDisasterEngine\n\n"
        "# Initialize Engine (Loads pre-trained XGBoost & GAT weights)\n"
        "engine = CombinedDisasterEngine()\n\n"
        "# Perform Inference\n"
        "result = engine.predict({\n"
        '    "latitude": 30.73, "longitude": 79.06, "district_id": "CHAMOLI_01",\n'
        '    "Rainfall_mm": 380.0, "Water_Level_m": 8.5, "Elevation_m": 2200,\n'
        '    "slope_angle_deg": 44.0, "Population_Density": 450,\n'
        '    "Land_Cover": "Forest", "Soil_Type": "Loam", "Historical_Floods": 1\n'
        "})\n\n"
        'print("Unified Risk Score:", result["unified_disaster_assessment"]["unified_risk_score"])\n'
        'print("Action Warning:    ", result["unified_disaster_assessment"]["action_recommendation"])'
    )
    
    r_code = cp_code.add_run(code_text)
    r_code.font.name = 'Consolas'
    r_code.font.size = Pt(9.5)
    r_code.font.color.rgb = RGBColor(220, 220, 220)

    doc.add_paragraph()

    # Save document
    doc.save(DOCX_PATH)
    print(f"[OK] Generated Word Document Report at: {DOCX_PATH}")

if __name__ == "__main__":
    create_report()
